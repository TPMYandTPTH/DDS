#!/usr/bin/env python3
"""
TP Document Divider System (DDS)
Flask Backend for Document Splitting
"""

import os
import io
import zipfile
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
import tempfile
import shutil
from datetime import datetime
from flask_cors import CORS

app = Flask(__name__)
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max

# Store processed files temporarily
processed_files = {}

def split_document_by_pages(input_path, pages_per_file=2):
    """Split a Word document into multiple documents with specified pages per file"""
    try:
        doc = Document(input_path)
        
        # Count page breaks
        page_breaks = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if '\f' in run.text or '\x0c' in run.text:
                    page_breaks += 1
        
        estimated_pages = max(page_breaks + 1, 1)
        
        # Create output directory
        output_dir = tempfile.mkdtemp()
        output_files = []
        
        # Initialize first document
        current_doc = Document()
        
        # Copy page settings
        for section in doc.sections:
            new_section = current_doc.sections[0]
            new_section.page_height = section.page_height
            new_section.page_width = section.page_width
            new_section.left_margin = section.left_margin
            new_section.right_margin = section.right_margin
            new_section.top_margin = section.top_margin
            new_section.bottom_margin = section.bottom_margin
            break
        
        page_count = 0
        file_count = 1
        
        # Process document elements
        for element in doc.element.body:
            if page_count >= pages_per_file:
                # Save current document
                output_path = os.path.join(output_dir, f'JD_{file_count:03d}.docx')
                current_doc.save(output_path)
                output_files.append(output_path)
                
                # Start new document
                current_doc = Document()
                for section in doc.sections:
                    new_section = current_doc.sections[0]
                    new_section.page_height = section.page_height
                    new_section.page_width = section.page_width
                    new_section.left_margin = section.left_margin
                    new_section.right_margin = section.right_margin
                    new_section.top_margin = section.top_margin
                    new_section.bottom_margin = section.bottom_margin
                    break
                
                page_count = 0
                file_count += 1
            
            # Add element to current document
            current_doc.element.body.append(element)
            
            # Check for page break
            if element.tag.endswith('p'):
                for run in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                    for br_elem in run.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'):
                        if br_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                            page_count += 1
        
        # Save last document
        if len(current_doc.element.body) > 0:
            output_path = os.path.join(output_dir, f'JD_{file_count:03d}.docx')
            current_doc.save(output_path)
            output_files.append(output_path)
        
        return {
            'success': True,
            'output_files': output_files,
            'file_count': len(output_files),
            'total_pages': estimated_pages,
            'output_dir': output_dir
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

@app.route('/')
def index():
    """Serve the HTML file"""
    return render_template('index.html')

@app.route('/health')
def health():
    """Health check endpoint"""
    return jsonify({'status': 'healthy', 'timestamp': datetime.now().isoformat()})

@app.route('/split', methods=['POST'])
def split_document():
    """Handle document splitting"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'No file uploaded'})
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'success': False, 'message': 'No file selected'})
        
        if not (file.filename.endswith('.docx') or file.filename.endswith('.docm')):
            return jsonify({'success': False, 'message': 'Invalid file type. Please upload .docx or .docm'})
        
        # Save uploaded file
        temp_input = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        file.save(temp_input.name)
        temp_input.close()
        
        # Process document
        result = split_document_by_pages(temp_input.name, pages_per_file=2)
        
        # Clean up input
        os.unlink(temp_input.name)
        
        if result['success']:
            session_id = datetime.now().strftime('%Y%m%d_%H%M%S')
            processed_files[session_id] = result
            
            return jsonify({
                'success': True,
                'message': f'Successfully split document into {result["file_count"]} files',
                'file_count': result['file_count'],
                'total_pages': result['total_pages'],
                'session_id': session_id
            })
        else:
            return jsonify({
                'success': False,
                'message': f'Error: {result["error"]}'
            })
            
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Server error: {str(e)}'
        })

@app.route('/download/<session_id>')
def download_files(session_id):
    """Download all split files as ZIP"""
    try:
        if session_id not in processed_files:
            return jsonify({'success': False, 'message': 'Session expired'})
        
        result = processed_files[session_id]
        
        # Create ZIP in memory
        memory_file = io.BytesIO()
        with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
            for file_path in result['output_files']:
                file_name = os.path.basename(file_path)
                zf.write(file_path, file_name)
        
        memory_file.seek(0)
        
        # Clean up
        shutil.rmtree(result['output_dir'])
        del processed_files[session_id]
        
        return send_file(
            memory_file,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'TP_JobDescriptions_{session_id}.zip'
        )
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Download error: {str(e)}'})

if __name__ == '__main__':
    print("=" * 60)
    print("TP Document Divider System (DDS)")
    print("=" * 60)
    print("Server starting...")
    print("Access: http://localhost:5000")
    print("Press CTRL+C to stop")
    print("=" * 60)
    app.run(debug=True, host='0.0.0.0', port=5000)
